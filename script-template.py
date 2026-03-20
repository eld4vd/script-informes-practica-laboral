"""
Generador de Informe Semanal de Actividades - Práctica Laboral
Plantilla Base - Versión Pública para GitHub

Uso:
    python script-template.py

Edita las variables en la sección "DATOS A COMPLETAR" antes de ejecutar.
Requiere: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ============================================================
#  DATOS A COMPLETAR - Edita aquí cada semana
# ============================================================

NUMERO_INFORME = "[NUMERO]"
PERIODO_SEMANAL = "[DIA/MES/AÑO] - [DIA/MES/AÑO]"

PRACTICANTE = "[NOMBRE DEL PRACTICANTE]"
DOCENTE_PRACTICA = "[NOMBRE DEL DOCENTE]"
ENTIDAD = "[NOMBRE DE LA ENTIDAD]"
AREA_INTERNA = "[ÁREA O DEPARTAMENTO]"
RESPONSABLE_AREA = "[NOMBRE DEL RESPONSABLE]"
FECHA_INICIO = "[FECHA INICIO]"
FECHA_FIN = "[FECHA FIN]"

# Jornadas: lista de 5 dicts (una por día laboral)
JORNADAS = [
    {
        "numero": "1",
        "fecha": "[FECHA 1]",
        "hora_ingreso": "08:00",
        "hora_salida": "12:00",
        "actividades": "[ACTIVIDAD DÍA 1]",          
        "habilidades": "[HABILIDAD DÍA 1]",          
        "dificultades": "[DIFICULTAD DÍA 1]",         
        "resultados": "[RESULTADO DÍA 1]",           
        "observaciones": "[OBSERVACIÓN DÍA 1]",        
    },
    {
        "numero": "2",
        "fecha": "[FECHA 2]",
        "hora_ingreso": "08:00",
        "hora_salida": "12:00",
        "actividades": "[ACTIVIDAD DÍA 2]",
        "habilidades": "[HABILIDAD DÍA 2]",
        "dificultades": "[DIFICULTAD DÍA 2]",
        "resultados": "[RESULTADO DÍA 2]",
        "observaciones": "[OBSERVACIÓN DÍA 2]",
    },
    {
        "numero": "3",
        "fecha": "[FECHA 3]",
        "hora_ingreso": "08:00",
        "hora_salida": "12:00",
        "actividades": "[ACTIVIDAD DÍA 3]",
        "habilidades": "[HABILIDAD DÍA 3]",
        "dificultades": "[DIFICULTAD DÍA 3]",
        "resultados": "[RESULTADO DÍA 3]",
        "observaciones": "[OBSERVACIÓN DÍA 3]",
    },
    {
        "numero": "4",
        "fecha": "[FECHA 4]",
        "hora_ingreso": "08:00",
        "hora_salida": "12:00",
        "actividades": "[ACTIVIDAD DÍA 4]",
        "habilidades": "[HABILIDAD DÍA 4]",
        "dificultades": "[DIFICULTAD DÍA 4]",
        "resultados": "[RESULTADO DÍA 4]",
        "observaciones": "[OBSERVACIÓN DÍA 4]",
    },
    {
        "numero": "5",
        "fecha": "[FECHA 5]",
        "hora_ingreso": "08:00",
        "hora_salida": "12:00",
        "actividades": "[ACTIVIDAD DÍA 5]",
        "habilidades": "[HABILIDAD DÍA 5]",
        "dificultades": "[DIFICULTAD DÍA 5]",
        "resultados": "[RESULTADO DÍA 5]",
        "observaciones": "[OBSERVACIÓN DÍA 5]",
    },
]

FIRMA_PRACTICANTE = "[FIRMA PRACTICANTE]"
FIRMA_RESPONSABLE = "[FIRMA RESPONSABLE]"

NOMBRE_ARCHIVO_SALIDA = f"Informe_Semanal_{NUMERO_INFORME}.docx"

# ============================================================
#  Colores y fuentes del documento original
# ============================================================

COLOR_TEXTO     = RGBColor(0x53, 0x51, 0x51)   # #535151 gris oscuro
COLOR_BORDE     = "757171"                       # gris borde tablas
COLOR_FONDO_HDR = "DBDBDB"                       # gris encabezados
FUENTE          = "Ebrima"


# ============================================================
#  Helpers de formato
# ============================================================

def set_cell_border(cell, color=COLOR_BORDE):
    """Aplica bordes simples a una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_shading(cell, fill=COLOR_FONDO_HDR):
    """Aplica color de fondo a una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "000000")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def cell_para(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT,
              color=COLOR_TEXTO, italic=False):
    """Limpia la celda y escribe un párrafo con formato."""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for p in cell.paragraphs:
        p.clear()
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.font.name = FUENTE
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    # Espaciado
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)
    return para


def add_paragraph(doc, text, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.LEFT,
                  color=COLOR_TEXTO, space_before=0, space_after=0, shading=None):
    """Agrega un párrafo al documento con formato completo."""
    para = doc.add_paragraph()
    para.alignment = align
    pPr = para._p.get_or_add_pPr()

    if shading:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), shading)
        pPr.append(shd)

    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(space_before))
    sp.set(qn("w:after"), str(space_after))
    pPr.append(sp)

    run = para.add_run(text)
    run.font.name = FUENTE
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    return para


# ============================================================
#  Construcción del documento
# ============================================================

def build_doc():
    doc = Document()

    # Márgenes (1.5 cm aprox laterales, 2 cm arriba/abajo)
    section = doc.sections[0]
    section.top_margin    = Cm(2.5) # Aumentar un poco para el header
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    # --- ENCABEZADO ---
    header = section.header
    header.is_linked_to_previous = False
    
    # Limpiamos el párrafo por defecto
    for p in header.paragraphs:
        p.text = ""

    # El ancho del área de contenido en papel Carta (21.59cm) menos márgenes (2.5x2) = 16.59cm
    htable = header.add_table(rows=2, cols=2, width=Cm(16.6))
    htable.autofit = False
    
    for row in htable.rows:
        row.cells[0].width = Cm(8.3)
        row.cells[1].width = Cm(8.3)
    
    # Quitar bordes de la tabla de encabezado por defecto
    for row in htable.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for b in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                bnd = OxmlElement(f"w:{b}")
                bnd.set(qn("w:val"), "nil")
                tcBorders.append(bnd)
            tcPr.append(tcBorders)

    c0 = htable.cell(0, 0)
    c1 = htable.cell(0, 1)
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Celda 0,0: Logo a la izquierda
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    try:
        # Intenta cargar 'logo.png' si existe en la misma carpeta
        p0.add_run().add_picture("logo.png", width=Cm(1.1)) # Imagen más pequeña (1.1 cm)
    except Exception as e:
        run_logo = p0.add_run("[IMAGEN LOGO FALTANTE: Guarda tu imagen como 'logo.png' en esta misma carpeta]")
        run_logo.font.size = Pt(8)
        run_logo.font.color.rgb = RGBColor(255, 0, 0)

    # Celda 0,1: Texto de la sigla a la derecha
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run1 = p1.add_run("SHC170 - I/2026")
    run1.font.name = FUENTE
    run1.font.size = Pt(11) # Tamaño de fuente cambiado a 11
    run1.font.bold = True
    run1.font.color.rgb = RGBColor(0, 0, 0)

    # Celda inferior 1,0 fusionada para la franja gris
    c_merged = htable.cell(1, 0).merge(htable.cell(1, 1))
    
    # Aplicar el color de fondo gris a la celda
    set_cell_shading(c_merged, fill="DBDBDB") # Color gris de sistema
    
    # Para la franja gris (sin usar demasiados espacios para no deformar los márgenes):
    p_merged = c_merged.paragraphs[0]
    p_merged.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_espacios = p_merged.add_run(" ") # Un solo espacio, el color de la celda cubre todo el ancho
    run_espacios.font.name = FUENTE
    run_espacios.font.size = Pt(20) # Tamaño aumentado a 20 para hacer la franja gris más gruesa
    
    # Ajustamos espaciados para que el grosor se aplique bien
    pPr = p_merged._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:before"), "0")
    pPr.append(spacing)

    # Añadir un salto de línea/párrafo vacío extra en el encabezado 
    # para separar del contenido principal sin que sea gigante
    p_separador = header.add_paragraph()
    p_separador.paragraph_format.space_before = Pt(0)
    p_separador.paragraph_format.space_after = Pt(8) # Espacio normal, ya no es tan grande

    # --- Títulos principales ---
    add_paragraph(doc, "INFORME SEMANAL DE ACTIVIDADES",
                  bold=True, size=26, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "PRÁCTICA LABORAL",
                  bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER)

    # --- Sección 1: DATOS GENERALES ---
    add_paragraph(doc, "1. DATOS GENERALES",
                  bold=True, size=16, space_before=240, shading=COLOR_FONDO_HDR)
    add_paragraph(doc, "", size=8)  # espaciado pequeño

    # Tabla datos generales (4 columnas)
    tbl1 = doc.add_table(rows=5, cols=4)
    tbl1.style = "Table Grid"

    def set_datos_row_height(row, val="700"):
        trPr = row._tr.get_or_add_trPr()
        trHeight = OxmlElement("w:trHeight")
        trHeight.set(qn("w:val"), val)
        trPr.append(trHeight)

    for row in tbl1.rows:
        set_datos_row_height(row, "700")

    def fill_datos_row(row_idx, label1, val1, label2, val2,
                       span_val1=False, span_label2=False):
        row = tbl1.rows[row_idx]
        cells = row.cells

        # Celda etiqueta 1 (fondo gris)
        set_cell_shading(cells[0])
        set_cell_border(cells[0])
        cell_para(cells[0], label1, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Celda valor 1
        set_cell_border(cells[1])
        cell_para(cells[1], val1, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

        if label2 is not None:
            set_cell_shading(cells[2])
            set_cell_border(cells[2])
            cell_para(cells[2], label2, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_border(cells[3])
            cell_para(cells[3], val2, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            # Fusionar celdas 2 y 3 cuando no hay segunda columna
            cells[1].merge(cells[3])
            cell_para(cells[1], val1, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    fill_datos_row(0, "No. Informe:", NUMERO_INFORME,
                      "Periodo Semanal:", PERIODO_SEMANAL)
    fill_datos_row(1, "Practicante:", PRACTICANTE,
                      "Docente Práctica Laboral:", DOCENTE_PRACTICA)

    # Fila Entidad (ocupa columnas 1-3 fusionadas)
    row2 = tbl1.rows[2]
    set_cell_shading(row2.cells[0])
    set_cell_border(row2.cells[0])
    cell_para(row2.cells[0], "Nombre Entidad / Institución:", bold=True, size=9,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    merged = row2.cells[1].merge(row2.cells[2]).merge(row2.cells[3])
    set_cell_border(merged)
    cell_para(merged, ENTIDAD, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    fill_datos_row(3, "Área Interna:", AREA_INTERNA,
                      "Nombre Responsable Área Entidad:", RESPONSABLE_AREA)
    fill_datos_row(4, "Fecha de Inicio de Práctica:", FECHA_INICIO,
                      "Fecha de finalización de Práctica:", FECHA_FIN)

    add_paragraph(doc, "", size=6)

    # --- Sección 2: DETALLES DE ACTIVIDADES ---
    add_paragraph(doc, "2. DETALLES DE ACTIVIDADES REALIZADAS",
                  bold=True, size=16, space_before=240, shading=COLOR_FONDO_HDR)

    for j in JORNADAS:
        add_paragraph(doc, "", size=6)

        # Tabla de jornada: 4 columnas lógicas
        # Col 0: Actividades | Col 1: Habilidades | Col 2: Dificultades (fusionada x2) | Col 3: Resultados (fusionada x2)
        # Usamos 6 columnas físicas para poder hacer el encabezado con hora ingreso/salida
        tbl = doc.add_table(rows=4, cols=6)
        tbl.style = "Table Grid"

        def set_row_height(row, val="700"):
            trPr = row._tr.get_or_add_trPr()
            trHeight = OxmlElement("w:trHeight")
            trHeight.set(qn("w:val"), val)
            trPr.append(trHeight)

        # --- Fila 0: encabezado con fecha y horario ---
        r0 = tbl.rows[0].cells
        set_row_height(tbl.rows[0], "560")

        set_cell_shading(r0[0])
        set_cell_border(r0[0])
        cell_para(r0[0], f"Jornada {j['numero']}:", bold=True, size=9,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

        set_cell_border(r0[1])
        cell_para(r0[1], j["fecha"], size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

        set_cell_shading(r0[2])
        set_cell_border(r0[2])
        cell_para(r0[2], "Hora Ingreso:", bold=True, size=9,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

        set_cell_border(r0[3])
        cell_para(r0[3], j["hora_ingreso"], bold=True, size=9,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

        set_cell_shading(r0[4])
        set_cell_border(r0[4])
        cell_para(r0[4], "Hora Salida:", bold=True, size=9,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

        set_cell_border(r0[5])
        cell_para(r0[5], j["hora_salida"], bold=True, size=9,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

        # --- Fila 1: cabeceras de contenido (4 columnas lógicas) ---
        # Actividades | Habilidades | Dificultades(cols 2+3) | Resultados(cols 4+5)
        r1 = tbl.rows[1].cells
        set_row_height(tbl.rows[1], "700")

        set_cell_shading(r1[0])
        set_cell_border(r1[0])
        cell_para(r1[0], "Actividades", bold=True, size=9,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

        set_cell_shading(r1[1])
        set_cell_border(r1[1])
        cell_para(r1[1], "Habilidades aprendidas y/o utilizadas", bold=True, size=9,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

        # Dificultades: fusionar cols 2 y 3
        merged_dif_hdr = r1[2].merge(r1[3])
        set_cell_shading(merged_dif_hdr)
        set_cell_border(merged_dif_hdr)
        cell_para(merged_dif_hdr, "Dificultades presentadas / superadas", bold=True, size=9,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

        # Resultados: fusionar cols 4 y 5
        merged_res_hdr = r1[4].merge(r1[5])
        set_cell_shading(merged_res_hdr)
        set_cell_border(merged_res_hdr)
        cell_para(merged_res_hdr, "Resultados alcanzados", bold=True, size=9,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

        # --- Fila 2: contenido (misma estructura fusionada) ---
        r2 = tbl.rows[2].cells
        set_row_height(tbl.rows[2], "1500")

        set_cell_border(r2[0])
        cell_para(r2[0], j["actividades"], size=9)

        set_cell_border(r2[1])
        cell_para(r2[1], j["habilidades"], size=9)

        merged_dif = r2[2].merge(r2[3])
        set_cell_border(merged_dif)
        cell_para(merged_dif, j["dificultades"], size=9)

        merged_res = r2[4].merge(r2[5])
        set_cell_border(merged_res)
        cell_para(merged_res, j["resultados"], size=9)

        # --- Fila 3: encabezado "Observaciones Generales" ---
        r3 = tbl.rows[3].cells
        set_row_height(tbl.rows[3], "560")

        merged_obs_hdr = r3[0].merge(r3[1]).merge(r3[2]).merge(r3[3]).merge(r3[4]).merge(r3[5])
        set_cell_shading(merged_obs_hdr)
        set_cell_border(merged_obs_hdr)
        cell_para(merged_obs_hdr, "Observaciones Generales:", bold=True, size=9)

        # --- Fila 4: contenido de observaciones (fila extra) ---
        # Añadir via python-docx add_row
        obs_tr = tbl.add_row()
        set_row_height(obs_tr, "1200")
        obs_cells = obs_tr.cells
        merged_obs_val = obs_cells[0].merge(obs_cells[1]).merge(obs_cells[2]).merge(obs_cells[3]).merge(obs_cells[4]).merge(obs_cells[5])
        set_cell_border(merged_obs_val)
        cell_para(merged_obs_val, j["observaciones"], size=9)

    # --- Sección firmas ---
    # Añadimos más párrafos vacíos para empujar las firmas más abajo
    add_paragraph(doc, "", size=11)
    add_paragraph(doc, "", size=11)
    add_paragraph(doc, "", size=11)
    add_paragraph(doc, "", size=11)
    add_paragraph(doc, "", size=11)
    add_paragraph(doc, "", size=11)

    # Crear una tabla invisible de 2x2 para las firmas (nombres arriba, roles abajo)
    tbl_firmas = doc.add_table(rows=2, cols=2)
    tbl_firmas.autofit = False
    
    # Ancho total igual al de contenido (aprox 16.6 cm repartidos en dos)
    # Dejar un buen espacio entre ellos
    for row in tbl_firmas.rows:
        row.cells[0].width = Cm(8.3)
        row.cells[1].width = Cm(8.3)

    # Quitar todos los bordes
    for row in tbl_firmas.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for b in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                bnd = OxmlElement(f"w:{b}")
                bnd.set(qn("w:val"), "nil")
                tcBorders.append(bnd)
            tcPr.append(tcBorders)

    # Fila de nombres
    c0 = tbl_firmas.cell(0, 0)
    c1 = tbl_firmas.cell(0, 1)
    
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run0 = p0.add_run(FIRMA_PRACTICANTE)
    run0.font.name = FUENTE
    run0.font.size = Pt(11)
    run0.font.color.rgb = RGBColor(0, 0, 0)

    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run(FIRMA_RESPONSABLE)
    run1.font.name = FUENTE
    run1.font.size = Pt(11)
    run1.font.color.rgb = RGBColor(0, 0, 0)

    # Fila de roles
    c2 = tbl_firmas.cell(1, 0)
    c3 = tbl_firmas.cell(1, 1)

    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("PRACTICANTE")
    run2.font.name = FUENTE
    run2.font.size = Pt(11)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(0, 0, 0)

    p3 = c3.paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("RESPONSABLE PRÁCTICA LABORAL")
    run3.font.name = FUENTE
    run3.font.size = Pt(11)
    run3.font.bold = True
    run3.font.color.rgb = RGBColor(0, 0, 0)

    # Agregar unos saltos extra al final si hace falta
    add_paragraph(doc, "", size=11)

    doc.save(NOMBRE_ARCHIVO_SALIDA)
    print(f"✅ Documento generado: {NOMBRE_ARCHIVO_SALIDA}")


if __name__ == "__main__":
    build_doc()
